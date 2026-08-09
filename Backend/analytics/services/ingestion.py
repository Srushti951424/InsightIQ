"""Read uploaded files (csv/xlsx/xls/docx) into a single cleaned DataFrame,
and compute a data-quality report along the way.
"""
import io

import numpy as np
import pandas as pd

from .column_detection import detect_columns


class UnreadableFileError(Exception):
    pass


def _read_single_file(django_file) -> pd.DataFrame:
    name = django_file.name.lower()
    raw = django_file.read()

    if name.endswith('.csv'):
        return pd.read_csv(io.BytesIO(raw))
    if name.endswith('.xlsx') or name.endswith('.xls'):
        return pd.read_excel(io.BytesIO(raw))
    if name.endswith('.docx'):
        return _read_docx_tables(raw)
    if name.endswith('.doc'):
        raise UnreadableFileError(
            f'"{django_file.name}" is a legacy .doc file. Please save it as .docx or .csv and re-upload.'
        )
    raise UnreadableFileError(f'Unsupported file type: "{django_file.name}"')


def _read_docx_tables(raw_bytes) -> pd.DataFrame:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(raw_bytes))
    if not document.tables:
        raise UnreadableFileError('No tables found in the Word document.')

    frames = []
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) < 2:
            continue
        header, *body = rows
        frames.append(pd.DataFrame(body, columns=header))

    if not frames:
        raise UnreadableFileError('No usable tables found in the Word document.')
    return pd.concat(frames, ignore_index=True, sort=False)


def load_and_merge(files) -> tuple[pd.DataFrame, list[dict]]:
    """Read every uploaded file and concatenate them into one DataFrame.

    Returns (merged_df, per_file_meta) where per_file_meta mirrors what the
    frontend expects: [{name, size, rows}, ...]
    """
    frames = []
    meta = []
    for f in files:
        df = _read_single_file(f)
        frames.append(df)
        meta.append({'name': f.name, 'size': f.size, 'rows': int(len(df))})

    merged = pd.concat(frames, ignore_index=True, sort=False) if len(frames) > 1 else frames[0]
    return merged, meta


def clean_dataframe(df: pd.DataFrame) -> tuple[pd.DataFrame, list[dict], int]:
    """Clean a merged DataFrame and report what was done.

    Returns (cleaned_df, issues, quality_score).
    """
    issues = []
    original_rows = len(df)

    # Drop fully-empty rows/columns first (common artifact of Excel exports)
    df = df.dropna(axis=0, how='all')
    df = df.dropna(axis=1, how='all')

    # Normalize column names (strip whitespace) without renaming semantics
    df.columns = [str(c).strip() for c in df.columns]

    # De-duplicate exact-duplicate rows
    dup_count = int(df.duplicated().sum())
    if dup_count > 0:
        df = df.drop_duplicates()
        issues.append({
            'type': 'duplicate',
            'label': f'{dup_count} duplicate row{"s" if dup_count != 1 else ""} removed',
            'severity': 'low' if dup_count / max(original_rows, 1) < 0.05 else 'medium',
        })

    # Detect column roles on the deduplicated data
    detection = detect_columns(df)

    # Normalize the primary date column to ISO-8601 if one was found
    if detection['date_col']:
        col = detection['date_col']
        before_non_null = df[col].notna().sum()
        parsed = pd.to_datetime(df[col].astype(str), errors='coerce', format='mixed')
        if parsed.notna().sum() < before_non_null:
            try:
                parsed = pd.to_datetime(df[col].astype(str), errors='coerce')
            except Exception:
                pass
        df[col] = parsed
        if before_non_null > 0:
            issues.append({
                'type': 'format',
                'label': f'Column "{col}" normalized to a consistent date format',
                'severity': 'low',
            })

    # Report missing-value hotspots (>2% missing) without silently imputing,
    # so users know what's incomplete in their source data.
    missing_pct = (df.isna().mean() * 100).round(1)
    worst = missing_pct[missing_pct > 2].sort_values(ascending=False)
    for col, pct in list(worst.items())[:3]:
        issues.append({
            'type': 'missing',
            'label': f'{pct}% missing values in "{col}"',
            'severity': 'low' if pct < 15 else 'medium',
        })

    # Quality score: start at 100, subtract for duplicates removed, overall
    # missingness, and unreadable/near-empty columns.
    overall_missing = float(df.isna().mean().mean() * 100) if df.size else 0
    dup_penalty = min(15, (dup_count / max(original_rows, 1)) * 100)
    missing_penalty = min(30, overall_missing)
    score = int(round(max(0, 100 - dup_penalty - missing_penalty)))

    return df, issues, score
